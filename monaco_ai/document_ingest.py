from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup

from .config import Settings
from .utils import chunk_text, normalize_ws, sha256_bytes, sha256_text

DOCUMENT_EXTENSIONS = (".pdf", ".epub", ".mobi", ".azw3", ".djvu", ".docx", ".rtf", ".txt", ".md")
EBOOK_EXTENSIONS = (".pdf", ".epub", ".mobi", ".azw3", ".djvu")


@dataclass(slots=True)
class ExtractedDocument:
    title: str
    url: str
    text: str
    content_type: str = ""
    bytes_len: int = 0
    content_hash: str = ""
    text_hash: str = ""


def is_document_url(url: str) -> bool:
    """Return True only for real downloadable documents/e-books.

    HTML is deliberately NOT treated as a document. Research workers may crawl
    HTML pages to discover links, but they should not store every crawled page
    when the user explicitly asked for e-books. This prevents useless storage
    bloat from menus, tag pages and category pages.
    """
    path = urlparse(url).path.lower()
    return any(path.endswith(ext) for ext in DOCUMENT_EXTENSIONS)


def is_ebook_url(url: str) -> bool:
    path = urlparse(url).path.lower()
    return any(path.endswith(ext) for ext in EBOOK_EXTENSIONS)


def document_link_score(url: str, text: str = "") -> int:
    """Simple heuristic used by crawlers to prefer actual e-book/document links."""
    hay = f"{url} {text}".lower()
    score = 0
    if is_document_url(url):
        score += 100
    if is_ebook_url(url):
        score += 40
    for word in ("ebook", "e-book", "book", "pdf", "epub", "download", "manual", "whitepaper", "paper", "guide"):
        if word in hay:
            score += 8
    for bad in ("login", "signup", "cart", "basket", "share", "facebook", "twitter", "privacy", "terms"):
        if bad in hay:
            score -= 12
    return score


def clean_url(url: str) -> str:
    return url.split("#", 1)[0].strip()


def discover_links(html: str, base_url: str, same_domain_only: bool = True) -> list[str]:
    base_domain = urlparse(base_url).netloc.lower()
    soup = BeautifulSoup(html or "", "lxml")
    links: list[str] = []
    seen: set[str] = set()
    for a in soup.find_all("a", href=True):
        href = str(a.get("href") or "").strip()
        if not href or href.startswith(("#", "mailto:", "tel:", "javascript:")):
            continue
        full = clean_url(urljoin(base_url, href))
        parsed = urlparse(full)
        if parsed.scheme not in {"http", "https"}:
            continue
        if same_domain_only and parsed.netloc.lower() != base_domain:
            continue
        if full not in seen:
            seen.add(full)
            links.append(full)
    return links


def fetch_bytes(settings: Settings, url: str, timeout: int = 30, max_bytes: int = 80_000_000) -> tuple[bytes, str]:
    headers = {"User-Agent": settings.web_user_agent}
    r = requests.get(url, headers=headers, timeout=timeout, stream=True)
    r.raise_for_status()
    ctype = r.headers.get("content-type", "")
    chunks: list[bytes] = []
    total = 0
    for chunk in r.iter_content(1024 * 256):
        if not chunk:
            continue
        total += len(chunk)
        if total > max_bytes:
            raise ValueError(f"Document too large for safe ingest ({total} bytes > {max_bytes})")
        chunks.append(chunk)
    return b"".join(chunks), ctype


def extract_text_from_bytes(url: str, data: bytes, content_type: str = "") -> str:
    lower = urlparse(url).path.lower()
    if lower.endswith(".pdf") or "pdf" in content_type.lower():
        try:
            from pypdf import PdfReader  # type: ignore
            reader = PdfReader(io.BytesIO(data))
            parts = []
            for page in reader.pages[:500]:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    continue
            return normalize_ws("\n".join(parts))
        except Exception as exc:
            return f"[PDF extract failed: {type(exc).__name__}: {exc}]"
    if lower.endswith(".docx") or "word" in content_type.lower():
        try:
            from docx import Document  # type: ignore
            doc = Document(io.BytesIO(data))
            return normalize_ws("\n".join(p.text for p in doc.paragraphs if p.text))
        except Exception as exc:
            return f"[DOCX extract failed: {type(exc).__name__}: {exc}]"
    if lower.endswith((".mobi", ".azw3", ".djvu")):
        return ""
    if lower.endswith(".epub"):
        # Lightweight EPUB fallback: read XHTML/HTML items from the ZIP. No DRM bypass.
        try:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                parts = []
                for name in zf.namelist():
                    if name.lower().endswith((".html", ".xhtml", ".htm", ".txt")):
                        raw = zf.read(name)
                        txt = raw.decode("utf-8", errors="ignore")
                        soup = BeautifulSoup(txt, "lxml")
                        for tag in soup(["script", "style", "noscript", "svg"]):
                            tag.decompose()
                        parts.append(soup.get_text("\n"))
                return normalize_ws("\n".join(parts))
        except Exception as exc:
            return f"[EPUB extract failed: {type(exc).__name__}: {exc}]"
    text = data.decode("utf-8", errors="ignore")
    if "html" in content_type.lower() or lower.endswith((".html", ".htm")):
        soup = BeautifulSoup(text, "lxml")
        for tag in soup(["script", "style", "noscript", "svg"]):
            tag.decompose()
        return normalize_ws(soup.get_text("\n"))
    return normalize_ws(text)


def extract_title(url: str, text: str, content_type: str = "") -> str:
    path = Path(urlparse(url).path)
    if path.name:
        return path.name[:180]
    first = next((line.strip() for line in text.splitlines() if line.strip()), "")
    return first[:120] or url


def fetch_and_extract(settings: Settings, url: str) -> ExtractedDocument:
    data, ctype = fetch_bytes(settings, url)
    text = extract_text_from_bytes(url, data, ctype)
    title = extract_title(url, text, ctype)
    return ExtractedDocument(
        title=title,
        url=url,
        text=text,
        content_type=ctype,
        bytes_len=len(data),
        content_hash=sha256_bytes(data),
        text_hash=sha256_text(text or ""),
    )
